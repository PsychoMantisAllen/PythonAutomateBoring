import docx

d = docx.Document()  # select the file

print(d.paragraphs)

print(d.paragraphs[0].text)

p = d.paragraphs[1]
print(p.runs)
# run means the text in different format
# when the format change, the breakpoint of the change is one run
print(p.runs[0].text)

# we can even check if the run is bold, italic etc
print(p.runs[1].bold)

p.runs[3].text = 'haha what are we doing'  # change text

print(p.style)  # we can check the style
p.style = 'Title'  # and change it

d = docx.Document()  # create a new file
d.add_paragraph('Hello there')
d.save()
p = d.paragraphs[0]
p.add_run('This is a new run.')
p.runs[1].bold = True  # change the new run into bold


def gettext(filename):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        fullText.append(para.text)
    return '\n'.join(fullText)


print(gettext())
